import io
import json
from docx import Document
from fpdf import FPDF
from bs4 import BeautifulSoup

def generate_docx_bytes(content_json):
    """
    Given a dictionary of sections and content,
    generate a DOCX and return bytes.
    """
    if isinstance(content_json, str):
        try:
            content_json = json.loads(content_json)
        except:
            content_json = {"Content": content_json}
            
    doc = Document()
    for section, content in content_json.items():
        doc.add_heading(section.title(), level=1)
        
        # Unescape and strip HTML
        if isinstance(content, str):
            content_text = content.encode("utf-8").decode("unicode_escape")
            soup = BeautifulSoup(content_text, "html.parser")
            clean_text = soup.get_text()
        else:
            clean_text = str(content)
            
        doc.add_paragraph(clean_text)
        
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

def generate_pdf_bytes(content_json):
    """
    Given a dictionary of sections and content,
    generate a PDF and return bytes.
    """
    if isinstance(content_json, str):
        try:
            content_json = json.loads(content_json)
        except:
            content_json = {"Content": content_json}
            
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Use built-in font
    pdf.set_font("Helvetica", size=12)
    
    for section, content in content_json.items():
        pdf.set_font("Helvetica", style='B', size=16)
        # Avoid charmap errors by ignoring/replacing non-latin1 characters
        safe_section = section.title().encode('latin-1', 'replace').decode('latin-1')
        pdf.cell(200, 10, txt=safe_section, new_x="LMARGIN", new_y="NEXT")
        
        pdf.set_font("Helvetica", size=12)
        
        if isinstance(content, str):
            content_text = content.encode("utf-8").decode("unicode_escape")
            soup = BeautifulSoup(content_text, "html.parser")
            clean_text = soup.get_text()
        else:
            clean_text = str(content)
            
        safe_text = clean_text.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, txt=safe_text)
        pdf.ln(5)
        
    # pdf.output() returns a bytearray in fpdf2, which Streamlit might not like. Convert to bytes.
    return bytes(pdf.output())
